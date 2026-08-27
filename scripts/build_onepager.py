from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "NguyenQuangHuy_Day25_onepager.docx"

NAVY = "183B56"
TEAL = "0F766E"
BLUE_GRAY = "E8EEF5"
LIGHT = "F4F8FB"
MUTED = "5B6770"
INK = "1F2937"
GOLD = "7A5A00"
ORANGE = "FCE8D5"
RED = "9B1C1C"
GREEN = "1F7A4C"
WHITE = "FFFFFF"

# Named override: one-page memo density. The missing official template means this is a
# transparent, readable executive brief rather than an imitation of an unavailable file.
PAGE_MARGIN = Inches(0.45)
CONTENT_DXA = 10944
TABLE_INDENT = 0


def set_font(run, size=8.3, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=65, start=90, bottom=65, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_table_borders(table, color="D9E2EC", sz="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), sz)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_text(paragraph, text, size=8.3, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def paragraph(cell, text="", size=8.3, color=INK, bold=False, italic=False, after=2, before=0, align=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    if align is not None:
        p.alignment = align
    add_text(p, text, size=size, color=color, bold=bold, italic=italic)
    return p


def labeled(cell, label, text, size=8.0, after=2, color=INK):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    add_text(p, f"{label}: ", size=size, color=NAVY, bold=True)
    add_text(p, text, size=size, color=color)
    return p


def block_title(cell, title, subtitle):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    add_text(p, title, size=11.1, color=TEAL, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.line_spacing = 1.0
    add_text(p2, subtitle, size=7.5, color=MUTED, italic=True)


def metric_table(doc):
    table = doc.add_table(rows=1, cols=6)
    set_table_geometry(table, [1824] * 6)
    set_table_borders(table, color="C9D6E2", sz="5")
    values = [
        ("7,648", "VND / JOB", NAVY),
        ("22,945", "PRICE FLOOR", NAVY),
        ("32,000 + 30M", "USAGE + PLATFORM", TEAL),
        ("82.8%", "BLENDED GM", GREEN),
        ("80% EVAL", "PACKAGE COMPLETION", GOLD),
        ("44.3%", "BREAKEVEN PACKAGE RATE", RED),
    ]
    for cell, (value, label, accent) in zip(table.rows[0].cells, values):
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        add_text(p, value, size=11.5, color=accent, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.0
        add_text(p2, label, size=6.5, color=MUTED, bold=True)
    return table


def add_cell_note(cell, text, fill=ORANGE, color=GOLD):
    note = cell.add_table(rows=1, cols=1)
    set_table_geometry(note, [3480])
    set_table_borders(note, color="E9C8A0", sz="4")
    c = note.cell(0, 0)
    set_cell_shading(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    add_text(p, text, size=7.3, color=color, bold=True)
    cell.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(8.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.02

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "DAY25 | P-015 | PRICING + GTM + EVIDENCE", size=7.0, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    add_text(footer, "READY WITH DISCLOSED LIMITATIONS - see evidence/ and FINAL_AUDIT_REPORT.md | 2026-08-27", size=6.7, color=MUTED)

    # Title stack
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_text(p, "P-015 | AI FRAUD INVESTIGATION COPILOT", size=20, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    add_text(p, "Pricing, GTM and evidence memo | Nguyen Quang Huy | 2A202601873", size=9.0, color=MUTED, bold=True)

    lead = doc.add_table(rows=1, cols=1)
    set_table_geometry(lead, [CONTENT_DXA])
    set_table_borders(lead, color="B6D6D2", sz="6")
    c = lead.cell(0, 0)
    set_cell_shading(c, "E9F5F3")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    add_text(p, "WHAT IS SOLD", size=7.0, color=TEAL, bold=True)
    add_text(p, "  A grounded investigation package for each eligible suspicious alert, embedded into an existing fraud-operations workflow. The AI assembles evidence and policy context; an authorized human owns the final disposition.", size=8.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    metric_table(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # Three content blocks
    body = doc.add_table(rows=1, cols=3)
    set_table_geometry(body, [3560, 3560, 3560])
    set_table_borders(body, color="D9E2EC", sz="5")
    pricing, gtm, evidence = body.rows[0].cells
    for cell in (pricing, gtm, evidence):
        set_cell_shading(cell, "FFFFFF")
        # Remove the default empty paragraph's visible impact.
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    block_title(pricing, "01 | PRICING", "Sell the software workflow, not an unproven autonomous outcome")
    labeled(pricing, "Customer", "Mid-market fintech / payment processor")
    labeled(pricing, "Buyer / budget", "Head of Risk or COO / software and fraud-operations tooling")
    labeled(pricing, "One job", "An eligible alert is complete when the copilot returns a schema-valid, grounded investigation package with required evidence/decision fields; customer final disposition is recorded separately.")
    labeled(pricing, "Value metric", "Hybrid: 30M VND platform fee/month + 32,000 VND per completed job.")
    labeled(pricing, "Base fee", "30M/month purchases account-level workflow access, integration/configuration and governance/auditability; it is not a second charge for the package itself.", size=7.45)
    labeled(pricing, "Unit GM", "76.1% on variable usage; 82.8% blended after the 30M/month platform fee at the base case.", size=7.35)
    labeled(pricing, "Breakeven package-completion rate", "44.3% at the 60% usage-GM target; current commercial package completion is 80%, for a +35.7 pp package-completion buffer. Autonomous containment is 20% and is not compared with this threshold.", size=7.1)
    labeled(pricing, "Rationale", "Attribution 2/5: the 5-case local control eval produced 4/5 grounded packages (80%), but no customer time-saved or outcome measurement. Autonomy 1/5: autonomous containment is 1/5 (20%), while 4/5 cases require human review. Hybrid follows the low-attribution/low-autonomy matrix and avoids unsupported outcome billing.", size=7.4, after=3)
    labeled(pricing, "Value anchor", "15 min evidence assembly displaced at 250,000 VND/hour = 62,500 VND/job full value; 50% capture = 31,250 VND/job. Proposed 32,000 VND is above the 22,945 VND floor and below full labor value.", size=7.5)
    add_cell_note(pricing, "RED TEAM: if the 80% commercial package-completion rate is wrong by 2x to 40%, usage GM falls to 56.1%. Autonomous containment is 20% and is not compared with the package-completion threshold.")

    block_title(gtm, "02 | GTM", "One channel for 90 days: Sales-Led")
    labeled(gtm, "Affordability", "ARPU 106.8M VND/month | ACV 1.282B VND / $49.3K reference | 12-month payback", size=7.7)
    labeled(gtm, "CAC math", "Budget 1.061B | estimated CAC 350M | gap 711M | affordability ratio 3.03x", size=7.6)
    labeled(gtm, "Deals / AE", "2.81 deals/year | 0.013 deals/selling day; mathematically feasible but planning-only.", size=7.6)
    labeled(gtm, "Pain Moment", "09:00-11:00 after an overnight alert burst: a fraud analyst triages high-risk transactions in the existing fraud-operations alert queue/dashboard and needs a grounded case package before manual disposition.", size=7.15)
    labeled(gtm, "Surface", "User-facing: existing fraud-analyst alert investigation queue/dashboard side panel. Backend: Kafka fraud_alerts -> fraud engine/agent -> REST /api/v1/fraud/analyze. Live integration is PARTIAL / NOT VERIFIED.", size=7.05)
    labeled(gtm, "Month 1 | Learn", "Target 2 design partners, 8 interviews, 300 labeled jobs; produce pain notes, failure taxonomy and time-on-task baseline.", size=7.5)
    labeled(gtm, "Months 2-3 | Leverage", "Target 2 pilots / 6,000 jobs / >=85% package completion / autonomous containment measured separately / >=70% GM / CAC <=350M; track 12 qualified opportunities and 3 wins.", size=7.35)
    labeled(gtm, "Month 4+ | Expand", "Only after >=85% package completion, autonomous containment and 2 paying customers plus security-gap closure. Next niche: multi-merchant payment processors.", size=7.35)

    block_title(evidence, "03 | EVIDENCE", "Procurement can see what is proven and what is open")
    labeled(evidence, "Eval Results", "PARTIAL. Small local control eval: 5 attempted, commercial package completion 4/5 (80%), autonomous containment 1/5 (20%), human-review/escalation 4/5 (80%), grounding failures 1/5 (20%); retry telemetry is not instrumented. Eight RAG queries: 100% citation accuracy. Upstream ML recall: 80.08% validation / 84.84% integration. Not customer evidence.", size=7.0)
    labeled(evidence, "Risk Checklist", "PARTIAL. Grounding, fallback, injection defense, score preservation, redaction and human review are evidenced. Auth/RBAC, rate limit, TLS, retention/deletion, durable audit export, vendor terms and live deployment remain open.", size=7.35)
    labeled(evidence, "Pilot Report", "MISSING EVIDENCE. No customer or participant is fabricated. PLANNED - NOT YET EXECUTED: start 2026-10-01, end 2026-11-26, 2 design partners, 6,000 jobs; owner Nguyen Quang Huy; measure completion, time saved, latency, retries, Cost/Job and safety.", size=7.25)
    labeled(evidence, "Submission gate", "READY WITH DISCLOSED LIMITATIONS. Official Day28 templates were not found locally; this memo and workbook are transparent substitutes. All numeric memo claims trace to workbook cells in evidence/onepager-traceability.md. Production readiness remains gated.", size=7.2)
    add_cell_note(evidence, "Decision: READY for artifact review/submission. Production claims remain gated on the planned pilot and open controls.", fill="FDECEC", color=RED)

    # Compact source line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_text(p, "Source trail: P-015 artifacts/agent + artifacts/ml/lightgbm; P-015 docs/business_product_direction.md, ai_fraud_agent.md, security.md, deployment_guide.md; current OpenAI, Supabase, Stripe Radar and Fingerprint pricing checked 2026-08-27. See workbook and evidence pack for full notes.", size=6.8, color=MUTED, italic=True)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
