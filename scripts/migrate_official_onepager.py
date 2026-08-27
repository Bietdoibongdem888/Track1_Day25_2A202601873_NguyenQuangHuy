from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "templates" / "official" / "Day25-AI-Product-GTM-One-Pager-Template.docx"
OUTPUT = ROOT / "deliverables" / "NguyenQuangHuy_Day25_onepager.docx"


def replace_paragraph(paragraph, text: str, size: float | None = None, color: str | None = None):
    """Replace paragraph text while preserving the official paragraph geometry/style."""
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", "Arial")
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def replace_cell(cell, text: str, size: float = 7.3, color: str | None = None):
    """Use the existing first paragraph/cell while preserving table geometry."""
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text, size=size, color=color)
    for extra in cell.paragraphs[1:]:
        replace_paragraph(extra, "", size=size)
    return paragraph


def main() -> None:
    doc = Document(str(SOURCE))
    paragraphs = doc.paragraphs
    tables = doc.tables

    # Headings and narrative slots from the official one-pager. Existing styles,
    # spacing, and table topology remain intact.
    paragraph_values = {
        0: "DAY 25 · AI-IN-ACTION · TRACK 1",
        1: "MONETIZATION ONE-PAGER",
        2: "Tên / Nhóm: Nguyen Quang Huy | Sản phẩm: P-015 AI Fraud Investigation Copilot",
        3: "Value Metric đã chọn: HYBRID (official $/job scenario below) | Kênh đã chọn: Sales-Led | Ngày: 2026-08-27",
        4: "Every numeric claim below traces to the official Day25 workbook or named P-015 evidence; 20% is autonomous containment, not 80% package completion.",
        7: "Software / fraud-operations tooling; buyer: Head of Risk / COO; an authorized customer analyst owns final disposition.",
        8: "Commercial recommendation: Hybrid (platform/governance access + measured usage). Official workbook economics model the pure $/completed-autonomous-job usage leg.",
        9: "Customer value is a grounded investigation package that reduces analyst work while keeping the final disposition customer-owned.",
        11: "HYBRID: platform/governance access plus usage. Attribution score 2/10; autonomy score 1/10; no customer causal outcome evidence, and 4/5 local cases require human review.",
        12: "Chosen unit: one completed grounded investigation package, priced per completed autonomous job in the official cost/pricing case.",
        13: "Evidence: local controlled evaluation reports 4/5 package completion and 1/5 autonomous containment; those are distinct rates.",
        14: "Market reason: Intercom and Zendesk publish adjacent outcome/resolution anchors, but a direct P-015 outcome price is not yet defensible.",
        16: "Intercom Fin: $0.99 per resolution/outcome. Zendesk AI agents: as low as $1.50 per resolution. Checked 2026-08-27; these are adjacent benchmarks, not P-015 customer proof.",
        17: "Intercom source: https://www.intercom.com/help/en/articles/8205718-fin-ai-agent-outcomes",
        18: "Zendesk source: https://www.zendesk.com/service/ai/top-ai-agents/",
        20: "Official model: $0.5422 Cost/Job (~₫14,150); $1.6267 floor (~₫42,447); $1.75 price (~₫45,668); 69.0% GM; 15.5% breakeven autonomous containment; current 20.0%.",
        22: "Anchor: the official 70% labor ceiling is $1.75/job (B14/B16/B11); Intercom $0.99 and Zendesk $1.50 are adjacent anchors. Proposed $1.75 is a planning decision, not a market quote.",
        23: "Value and floor checks pass: $1.75 is above the 3x Cost/Job floor and within the labor/value guardrails in 2_Pricing.",
        24: "The official usage case is profitable at the current 20% autonomous containment; Hybrid access pricing remains a commercial recommendation for pilot discovery.",
        26: "If autonomous containment falls below about 12.4%, Gross Margin falls below 50% at the $1.75/job usage price (the official sensitivity table starts at 50%).",
        31: "Sales-Led only for 90 days; founder-led design-partner motion targets mid-market fintech/payment processors.",
        32: "ARPU $1,050/month; CAC budget $13,043.88; estimated CAC $12,000; budget coverage 1.09x (official B23 shows estimated/budget 0.92x); deals/AE/day 0.180.",
        33: "The $3,000/opportunity input is a founder-led planning assumption, not CRM evidence; replace it after pilot funnel data.",
        35: "ARPU $1,050/month; ACV $12,600/year; CAC budget $13,043.88; estimated CAC $12,000; B23 estimated/budget 0.92x; budget/estimated 1.09x.",
        37: "09:00–11:00 after an overnight alert burst: fraud analyst triages high-risk alerts in the existing fraud-operations alert queue/dashboard before manual disposition.",
        41: "Existing fraud-operations alert investigation queue/dashboard side panel. Backend: Kafka fraud_alerts → fraud engine/agent → REST /api/v1/fraud/analyze; live integration is PARTIAL / NOT VERIFIED.",
        46: "Missing evidence is labeled with owner and deadline; no customer or production claim is fabricated.",
        49: "Prepared for a two-minute outsider read; the buyer test is not yet run, so the question count remains explicitly untested.",
        50: "☐ Buyer, product, and unit can be stated in one sentence — prepared answer is above; outsider test not yet run.",
        51: "☐ Price, margin, and break point can be found in the official table — prepared answer is above; outsider test not yet run.",
        52: "☐ Channel and next proof step are explicit — prepared answer is above; outsider test not yet run.",
        53: "Số câu người đọc phải hỏi lại: NOT YET TESTED (target ≤3)",
        54: "AI-IN-ACTION · TRACK 1 · official template migration · evidence-led planning",
    }
    for index, value in paragraph_values.items():
        replace_paragraph(paragraphs[index], value)

    # Official key-number table: 7 rows x 3 columns.
    key_rows = [
        ("Cost/Job", "$0.5422/job (~₫14,150)", "1_Cost_Job!B66; VND B69"),
        ("Giá sàn", "$1.6267/job (~₫42,447)", "2_Pricing!B7"),
        ("Giá đề xuất", "$1.75/job (~₫45,668)", "2_Pricing!B19"),
        ("GM", "69.0%", "2_Pricing!B21"),
        ("Breakeven", "15.5% autonomous", "2_Pricing!B33"),
        ("Containment hiện tại", "20.0% autonomous (1/5)", "1_Cost_Job!B10; evidence/containment-eval.csv"),
    ]
    key_table = tables[0]
    for row_index, row in enumerate(key_rows, start=1):
        for col_index, value in enumerate(row):
            replace_cell(key_table.cell(row_index, col_index), value, size=7.1 if col_index != 1 else 7.5)

    # Official channel evidence table: 6 rows x 3 columns.
    channel_rows = [
        ("ARPU", "$1,050/month", "4_Channel_Fit!B5"),
        ("CAC budget", "$13,043.88/customer", "4_Channel_Fit!B9"),
        ("Deals/AE/day", "0.180", "4_Channel_Fit!B16"),
        ("Estimated CAC", "$12,000", "4_Channel_Fit!B22"),
        ("Estimated / budget", "0.92x", "4_Channel_Fit!B23; inverse coverage 1.09x"),
    ]
    channel_table = tables[1]
    for row_index, row in enumerate(channel_rows, start=1):
        for col_index, value in enumerate(row):
            replace_cell(channel_table.cell(row_index, col_index), value, size=7.0 if col_index != 1 else 7.4)

    # Official 90-day plan table: rows are Kênh, Mục tiêu số khách, Việc cụ thể,
    # KPI đo được, Ai chịu trách nhiệm; columns are the three official phases.
    plan_rows = [
        ("Kênh", "Sales-Led founder-led design-partner motion", "Sales-Led paid pilot motion", "Sales-Led expansion within adjacent payment-processor accounts"),
        ("Mục tiêu số khách", "2 design partners", "2 pilot customers / 6,000 labeled jobs", "2 paying customers + 1 adjacent segment"),
        ("Việc cụ thể", "8 fraud-ops interviews; produce pain notes and security-gap log", "Run 2 pilots; instrument completion, autonomy, retries and latency", "Standardize procurement pack; add multi-merchant processor playbook"),
        ("KPI đo được", "8 interviews, 2 design partners, 300 labeled jobs", ">=85% package completion; autonomy separate; >=70% blended GM; CAC <=$32k", "2 paying customers, no unsafe auto-action, auth/RBAC and retention evidence closed"),
        ("Ai chịu trách nhiệm", "Nguyen Quang Huy + 2 design-partner fraud leads", "Nguyen Quang Huy; customer fraud team owns final disposition", "Nguyen Quang Huy + customer champion"),
    ]
    plan_table = tables[2]
    for row_index, row in enumerate(plan_rows, start=1):
        for col_index, value in enumerate(row):
            replace_cell(plan_table.cell(row_index, col_index), value, size=6.2 if col_index else 6.6)

    # Official evidence table: preserve status column and make missing proof explicit.
    evidence_rows = [
        ("Eval Results (Day21–22)", "PARTIAL", "5-case local controlled eval: 80% commercial package completion, 20% autonomous containment, 80% human review, 20% grounding failure; not customer evidence.", "Nguyen Quang Huy · 2026-09-30"),
        ("Risk Checklist (Day24)", "PARTIAL", "Grounding/fallback/redaction evidenced; auth/RBAC, retention, TLS, vendor terms and live deployment remain open.", "Nguyen Quang Huy · before pilot / 2026-10-15"),
        ("Pilot Report", "PLANNED / NOT EXECUTED", "2 design partners, 6,000 labeled jobs and paired manual baseline; no customer result claimed.", "Nguyen Quang Huy · 2026-12-04"),
    ]
    evidence_table = tables[3]
    for row_index, row in enumerate(evidence_rows, start=1):
        for col_index, value in enumerate(row):
            replace_cell(evidence_table.cell(row_index, col_index), value, size=6.4 if col_index == 2 else 6.8)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"WROTE {OUTPUT}")
    print(f"OFFICIAL_SOURCE {SOURCE}")
    print(f"PARAGRAPHS {len(doc.paragraphs)} TABLES {len(doc.tables)}")


if __name__ == "__main__":
    main()
